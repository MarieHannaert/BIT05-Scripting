#!/usr/bin/python3
import os 
import fnmatch
import re
from docx import Document
print("================================================================================")
print("This script will parse info for your gene of interest from GFF annotation files.")
gene_name=input("Enter gene name (e.g. gyrB): ")
#gene_name= "gyrB"
print("================================================================================")
print("Getting list of folders in main folder ncbi_datasets ...")
location=(os.getcwd()+"/ncbi_datasets")
#print(location)
ncbidata_list = os.listdir(location)
print("Found 3 folders: {}".format(ncbidata_list))
document = Document()
document.add_heading('Results for gene name :{}'.format(gene_name), 0)
for i in range(len(ncbidata_list)):
    ncbidata = ncbidata_list[i]
    for root, dirs, files in os.walk(location+"/"+ncbidata):
        for filename in fnmatch.filter(files,'genomic.gff'):
            #print("Found {} in {}".format(filename, root))
            file_location=location+"/"+ncbidata+"/"+filename
            print("--------------------------------------------------------------------------------")
            print("Parsing GFF file {} :{}".format((i+1),file_location))
            fileObj = open(file_location, "r")
            line_counter=0
            for line in fileObj.readlines():#why readlines, because it has commentlines in the beginning 
                match = re.search(gene_name, line)
                line_counter += 1
                if match: 
                    split_line=line.split("\t")
                    #print(split_line)
                    print("Found {} on line {} where feature = {} | start = {} end = {}" .format(gene_name,line_counter, split_line[2], split_line[3], split_line[4]))
                    if split_line[2] == "CDS":
                        #print(split_line)
                        document.add_heading("{}".format(ncbidata))
                        last_line=split_line[8].split(";")
                        #print(last_line)
                        #for i in range(len(last_line)):
                            #n=last_line[i]
                        for match2 in last_line:
                            if "locus_tag" in match2:
                                document.add_paragraph(str(match2)) 
                            if "product" in match2:
                                document.add_paragraph(str(match2))
                            if "protein_id" in match2:
                                document.add_paragraph(str(match2))
                            #if i == (11 or 12 or 13):
                                #print(last_line[11])
                             #   document.add_paragraph(str(last_line[11]))
                                #print(last_line[12])
                              #  document.add_paragraph(str(last_line[12]))
                                #print(last_line[13])
                               # document.add_paragraph(str(last_line[13]))
document.save("{}.docx".format(gene_name))
