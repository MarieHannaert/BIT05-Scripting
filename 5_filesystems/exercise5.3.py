from docx import Document
document=Document()

import os
import re

fastqc_dir = "/media/sf_SF/BIT05-scripting/5_filesystems-writing-and-reading/fastqc/fastqc/"

for SRR_DIR in os.listdir(fastqc_dir):
    if SRR_DIR.endswith(".html"):
        continue

    location=fastqc_dir + SRR_DIR
    for filename in os.listdir(location):
        document.add_heading('Sample_{}'.format(SRR_DIR), 0)
        document.add_heading('Summary', level=1)
        if filename == "fastqc_data.txt":
            with open (location+"/"+filename) as f:
                match= re.search('Total Sequences.*',f.read())
                print(match.group(0))

        elif filename == "summary.txt":
            with open (location+"/"+filename) as f:
                match= re.search('.*Basic Statistics',f.read())
                A=match.group(0)
                gesplitste_A=A.split("\t")
                print(gesplitste_A[1]+" "+gesplitste_A[0])

            with open (location+"/"+filename) as f:
                match= re.search('.*Per base sequence quality',f.read())
                B=match.group(0)
                gesplitste_B=B.split("\t")
                print(gesplitste_B[1]+" "+gesplitste_B[0])
                
        from docx.shared import Cm
        document.add_picture(location+"/Images/per_base_quality.png", width=Cm(10))
        document.add_page_break()
document.save('fastqc_summary.docx')