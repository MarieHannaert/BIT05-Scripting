import pandas as pd
from docx import Document


document = Document()
document.add_heading('Output of script Panda openen csv', 0)
csv_iterator = pd.read_csv('output_interactionsweb.csv', iterator=True, chunksize=1000)
for chunk in csv_iterator:
    for index, row in chunk.iterrows():
        print(row)
        document.add_paragraph(str(row))
        

document.save('panda_openen_csv_script.docx')
#

