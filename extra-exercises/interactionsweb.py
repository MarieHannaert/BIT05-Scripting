#!/usr/bin/python3
import urllib.request
from bs4 import BeautifulSoup
from docx import Document
import csv
from openpyxl import Workbook
# Get content data using urllib and create soup object
url = "http://www.cazy.org/GH100_all.html"
response = urllib.request.urlopen(url)
content = response.read()
soup = BeautifulSoup(content, 'html.parser')
print(soup)
wb = Workbook()
# Active worksheet
ws = wb.active
ws.title = "interactionsweb output"
document = Document()
document.add_heading('Output of script interactionsweb', 0)
tag_table = soup.find("table", {"id":"pos_onglet"})
#print(tag_table)
Title=(tag_table.find("span", class_="tabulation"))
document.add_paragraph(str(Title))
(print(Title))
bactitle=Title.text
document.add_paragraph(str(bactitle))
print(bactitle)
    

document.add_page_break()
with open('output_interactionsweb.csv', mode='w') as interactions_file:
    pers_writer = csv.writer(interactions_file, 
    delimiter=',')
    table_headers=tag_table.find_all("tr", id="line_titre")
    print(table_headers)
    row=table_headers[1]
    cells = row.find_all("td")
    temp_list2=["row_nr"]
    
    for i in range(len(cells)):
        cell = cells[i]
        cleand_text2=cell.text.strip()
        print(cleand_text2)
        temp_list2.append(cleand_text2)


    #for cell in cells:
     #   cleand_text2=cell.text.strip()
      #  print(cleand_text2)
    # temp_list2.append(cleand_text2)
    pers_writer.writerow(temp_list2)
    ws.append(temp_list2)
    


    
    info=tag_table.find_all("tr", {"valign":"top"})
    #for secrow in info:
    for i in range(len(info)):
        secrow = info[i]
        spec_info=secrow.find_all("td")
        print(spec_info)
        temp_list=[i+1]
        for info_text in spec_info:
            cleand_text=info_text.text.strip()
            print(cleand_text)
            temp_list.append(cleand_text)
        pers_writer.writerow(temp_list)
        ws.append(temp_list)
interactions_file.close()
tag_table1 = soup.table
inhoud=tag_table1.find_all("tr")
wb.save("interactionsweb_output.xlsx")
for cell2 in inhoud:
    print((cell2.find("th", class_="thsum")).text)
    print((cell2.find("td", class_="tdsum")).text)
    document.add_paragraph(str((cell2.find("th", class_="thsum")).text))
    document.add_paragraph(str((cell2.find("td", class_="tdsum")).text))


document.save('Output_interactionsweb_script.docx')


